from docx import Document

def reverse_numbers_order(file_to_read='', file_to_save=''):
    """
    this method reverse numbers in docs file
    """
    if len(file_to_read) == 0 or len(file_to_save) == 0:
        return False
    doc = Document(file_to_read)
    for para in doc.paragraphs:
        new_text = ''
        count_num = 0
        for char in para.text:
            if str(char) in '1234567890':
                if count_num == 0:
                    new_text += str(char)
                    count_num += 1
                else:
                    new_text = new_text[:-1 * count_num] + str(char) + new_text[-1 * count_num:]
                    count_num += 1
            else:
                new_text += char
                count_num = 0

        para.text = new_text

    return doc.save(file_to_save)

def latin_to_arabic(file_to_read='', file_to_save=''):
    """
        this method changes Latin numbers in docs file to Arabic numbers
    """
    if len(file_to_read) == 0 or len(file_to_save) == 0:
        return False
    doc = Document(file_to_read)
    for paragraph in doc.paragraphs:
        arabic_num_text = ''
        for character in paragraph.text:
            new_character = character.replace('1', '۱').replace('2', '۲') \
                .replace('3', '۳').replace('4', '۴').replace('5', '۵').replace('6', '۶') \
                .replace('7', '۷').replace('8', '۸').replace('9', '۹').replace('0', '۰')
            arabic_num_text += new_character

        paragraph.text = arabic_num_text

    return doc.save(file_to_save)
